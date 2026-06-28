# -*- coding: utf-8 -*-
"""把 开题报告.md 转成对齐清华 202401 表的 Word .docx。
轻量 markdown 子集解析:标题/段落/表格/代码块/引用/内联粗体。"""
import re, sys
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = "/home/lizhaoxi/.kitty/workspaces/s_9205bdbc/kaiti/report/开题报告.md"
OUT = "/home/lizhaoxi/.kitty/workspaces/s_9205bdbc/kaiti/report/选题报告及论文工作计划表_202401.docx"

doc = Document()

# 页面 A4 + 页边距
sec = doc.sections[0]
sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
sec.left_margin = sec.right_margin = Cm(2.8)
sec.top_margin = sec.bottom_margin = Cm(2.5)

# 默认正文样式:宋体 小四(12pt)
normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(12)
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

def set_cjk(run, font="宋体"):
    run.font.name = "Times New Roman"
    r = run._element
    rPr = r.get_or_add_rPr()
    rf = rPr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts"); rPr.append(rf)
    rf.set(qn("w:eastAsia"), font)
    rf.set(qn("w:ascii"), "Times New Roman")
    rf.set(qn("w:hAnsi"), "Times New Roman")

def add_runs(p, text, base_font="宋体", bold=False, size=None, color=None):
    """处理内联 **粗体**"""
    for i, seg in enumerate(re.split(r"\*\*", text)):
        if seg == "":
            continue
        run = p.add_run(seg)
        run.bold = bold or (i % 2 == 1)
        if size: run.font.size = Pt(size)
        if color: run.font.color.rgb = color
        set_cjk(run, base_font)

def set_cell_border(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single"); e.set(qn("w:sz"), "4")
        e.set(qn("w:space"), "0"); e.set(qn("w:color"), "808080")
        borders.append(e)
    tcPr.append(borders)

# 表单大标题
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_runs(t, "清华大学硕士生(学术学位)选题报告及论文工作计划表", base_font="黑体", bold=True, size=16)
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_runs(sub, "(202401 版)", base_font="楷体", size=11)
doc.add_paragraph()

lines = open(SRC, encoding="utf-8").read().splitlines()

def is_table_sep(s):
    return bool(re.match(r"^\s*\|[\s:|-]+\|\s*$", s))

i = 0
n = len(lines)
while i < n:
    line = lines[i]
    s = line.rstrip()
    # 跳过 md 文档自身的一级标题(# ...)和首个引用说明
    if s.startswith("# "):
        i += 1; continue
    # 代码块
    if s.startswith("```"):
        i += 1; buf = []
        while i < n and not lines[i].startswith("```"):
            buf.append(lines[i]); i += 1
        i += 1
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        for j, cl in enumerate(buf):
            run = p.add_run(cl)
            run.font.name = "Consolas"; run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
            if j != len(buf) - 1:
                run.add_break()
        continue
    # 表格
    if s.startswith("|") and i + 1 < n and is_table_sep(lines[i + 1]):
        rows = []
        header = [c.strip() for c in s.strip().strip("|").split("|")]
        rows.append(header)
        i += 2
        while i < n and lines[i].lstrip().startswith("|"):
            rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
            i += 1
        ncol = max(len(r) for r in rows)
        tbl = doc.add_table(rows=len(rows), cols=ncol)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.autofit = True
        for ri, row in enumerate(rows):
            for ci in range(ncol):
                cell = tbl.cell(ri, ci)
                set_cell_border(cell)
                txt = row[ci] if ci < len(row) else ""
                cp = cell.paragraphs[0]
                add_runs(cp, txt, base_font="宋体", bold=(ri == 0), size=10.5)
        doc.add_paragraph()
        continue
    # 标题
    if s.startswith("## "):
        p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(4)
        add_runs(p, s[3:].strip(), base_font="黑体", bold=True, size=15)
        i += 1; continue
    if s.startswith("### "):
        p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
        add_runs(p, s[4:].strip(), base_font="黑体", bold=True, size=12.5)
        i += 1; continue
    # 引用
    if s.startswith(">"):
        p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.5)
        add_runs(p, s.lstrip("> ").strip(), base_font="楷体", size=10.5, color=RGBColor(0x55,0x55,0x55))
        i += 1; continue
    # 分隔线
    if re.match(r"^\s*---+\s*$", s):
        i += 1; continue
    # 列表项
    m = re.match(r"^(\s*)([-*]|\d+\.)\s+(.*)$", s)
    if m:
        p = doc.add_paragraph(style=None)
        p.paragraph_format.left_indent = Cm(0.74)
        add_runs(p, ("• " if m.group(2) in ("-","*") else m.group(2)+" ") + m.group(3), base_font="宋体")
        i += 1; continue
    # 空行
    if s.strip() == "":
        i += 1; continue
    # 普通段落
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.25
    add_runs(p, s.strip(), base_font="宋体")
    i += 1

# 签字栏
doc.add_paragraph()
for title in ["指导教师评语:", "选题报告会(专家组)意见:", "选题报告考核成绩:"]:
    p = doc.add_paragraph(); add_runs(p, title, base_font="黑体", bold=True, size=12)
    box = doc.add_table(rows=1, cols=1); set_cell_border(box.cell(0,0))
    box.cell(0,0).paragraphs[0].add_run("\n\n\n")
    sign = doc.add_paragraph(); sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_runs(sign, "签字:　　　　　　　日期:　　　年　　月　　日", base_font="宋体", size=10.5)
    doc.add_paragraph()
note = doc.add_paragraph()
add_runs(note, "本材料留院(系、所)业务办公室,供论文中期检查、论文抽查时参考。", base_font="楷体", size=9, color=RGBColor(0x77,0x77,0x77))

doc.save(OUT)
print("SAVED:", OUT)
