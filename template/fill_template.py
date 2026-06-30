# -*- coding: utf-8 -*-
"""在官方 .docx 模板上原位填写(保留格式)。"""
import copy
from docx import Document
from docx.oxml.ns import qn

SRC="/home/lizhaoxi/.kitty/workspaces/s_9205bdbc/kaiti/template/选题报告模板_202401.docx"
OUT="/home/lizhaoxi/.kitty/workspaces/s_9205bdbc/kaiti/template/选题报告_202401_李兆曦_已填写.docx"
d=Document(SRC)

def repl(p, old, new):
    full="".join(r.text for r in p.runs)
    if old not in full:
        return False
    full=full.replace(old,new)
    if p.runs:
        p.runs[0].text=full
        for r in p.runs[1:]:
            r.text=""
    else:
        p.add_run(full)
    return True

def fill_empty(p, text, ref_run):
    """把空段落填上 text,字体复制 ref_run。"""
    if p.runs:
        p.runs[0].text=text
        for r in p.runs[1:]:
            r.text=""
        tgt=p.runs[0]
    else:
        tgt=p.add_run(text)
    # 复制 rPr(字体)以保格式
    if ref_run is not None and ref_run._r.rPr is not None:
        newrpr=copy.deepcopy(ref_run._r.rPr)
        old=tgt._r.find(qn('w:rPr'))
        if old is not None: tgt._r.remove(old)
        tgt._r.insert(0,newrpr)

# ===== 封面 =====
P=d.paragraphs
repl(P[8], "研究生姓名", "研究生姓名　李兆曦")
repl(P[10], "院       系", "院       系　计算机科学与技术系")
repl(P[13], "学       科", "学       科　计算机科学与技术")
repl(P[16], "指 导 教 师", "指 导 教 师　徐恪")
repl(P[16], "专业技术职务", "专业技术职务　教授")
repl(P[26], "年         月          日", "2026 年    6    月          日")

# ===== 表0:论文题目/密级/学科交叉 =====
t0=d.tables[0]
c=t0.rows[0].cells[0]
repl(c.paragraphs[0], "论文题目：", "论文题目：基于大模型编码智能体的自进化 eBPF 网络入侵防御系统")
repl(c.paragraphs[2], "公开（  ）", "公开（ √ ）")
repl(c.paragraphs[5], "否（  ）", "否（ √ ）")   # 学科交叉:否

# ===== 表0:文献情况 + 主要文献 =====
cf=t0.rows[2].cells[0]
ref0=cf.paragraphs[0].runs[0] if cf.paragraphs[0].runs else None
repl(cf.paragraphs[1], "国内文献约             篇，国外文献约                   篇。",
     "国内文献约  8  篇,国外文献约  50  篇。")
refs=[
 "[1] FENG X, LI Z, et al. Off-Path TCP Exploits: PMTUD Breaks TCP Connection Isolation in IP Address Sharing Scenarios. ACM CCS, 2025.",
 "[2] FENG X, et al. Exploiting Cross-Layer Vulnerabilities: Off-Path Attacks on the TCP/IP Protocol Suite. Communications of the ACM, 2025.",
 "[3] YANG Y, ZHENG Y, et al. Kgent: Kernel Extensions Large Language Model Agent. ACM SIGCOMM Workshop on eBPF, 2024.",
 "[4] GERSHUNI E, et al. Simple and Precise Static Analysis of Untrusted Linux Kernel Extensions. PLDI, 2019.",
 "[5] XU Q, et al. Synthesizing Safe and Efficient Kernel Extensions for Packet Processing. ACM SIGCOMM, 2021.",
 "[6] ROMERA-PAREDES B, et al. Mathematical Discoveries from Program Search with LLMs. Nature, 2024.",
 "[7] NOVIKOV A, et al. AlphaEvolve: A Coding Agent for Scientific and Algorithmic Discovery. arXiv:2506.13131, 2025.",
 "[8] YANG J, et al. SWE-agent: Agent-Computer Interfaces Enable Automated Software Engineering. NeurIPS, 2024.",
]
empties=[p for p in cf.paragraphs[3:] if not p.text.strip()]
for i,line in enumerate(refs):
    if i<len(empties): fill_empty(empties[i], line, ref0)

# ===== 表1:实验条件/安全/项目类别/论文类型/工作内容 =====
t1=d.tables[1]
# 实验设备/条件
r0=t1.rows[0].cells[0]
repl(r0.paragraphs[0], "实验设备或其他研究条件落实情况：",
     "实验设备或其他研究条件落实情况：具备 Linux 内核与 eBPF 开发 / 测试环境、GPU 算力与大模型 API 访问;并以开源系统 PacketScope/Guarder 作为工程底座,实验条件已落实。")
# 实验安全(打√)
r1=t1.rows[1].cells[0]
repl(r1.paragraphs[1], "是 （  ）", "是 （√）")
repl(r1.paragraphs[2], "是 （  ）", "是 （√）")
repl(r1.paragraphs[3], "否 （  ）", "否 （√）")
# 项目类别
r2=t1.rows[2].cells[0]
ref2=r2.paragraphs[0].runs[0] if r2.paragraphs[0].runs else None
# 追加一个段落写值(保持label不动)
np=r2.add_paragraph(); np.add_run("国家项目（国家重点研发计划 / 国家自然科学基金资助）")
if ref2 is not None and ref2._r.rPr is not None:
    rr=np.runs[0]; nr=copy.deepcopy(ref2._r.rPr); rr._r.insert(0,nr)
# 论文类型:重写 5 个类型单元格为普通文本(修复 ①②④⑤ 渲染错误),并给 ③ 标√
def setcell(cell, text):
    p=cell.paragraphs[0]
    if p.runs:
        p.runs[0].text=text
        for r in p.runs[1:]:
            r.text=""
    else:
        p.add_run(text)
setcell(t1.rows[4].cells[1], "① 基础研究")
setcell(t1.rows[4].cells[3], "② 应用基础（理论）研究")
setcell(t1.rows[5].cells[1], "③ 应用研究　√")
setcell(t1.rows[5].cells[3], "④ 开发研究")
setcell(t1.rows[6].cells[1], "⑤ 其它（详细注明）")
# 工作内容及日程
r7=t1.rows[7].cells[0]
ref7=r7.paragraphs[0].runs[0] if r7.paragraphs[0].runs else None
plan=[
 "本课题拟构建一套由大模型编码智能体自动维护的内核 eBPF 防御系统。主要内容与日程安排如下:",
 "2026.07–09:构建系统与专用 harness 原型——打通智能体“生成 eBPF→编译→verifier 校验→沙箱测试→加载→读取遥测”的工具链与闭环;",
 "2026.10–12:完善“生成可用且安全 eBPF”的能力(生成质量与运行安全),完成闭环原型与初步实验;",
 "2027.01–03:面向 PMTUD/NAT/IPID 等真实协议攻击的防御生成与评测,与人工规则及现有工具对比,并接受中期检查;",
 "2027.04–05:论文撰写、预答辩与盲审送审;",
 "2027.06:学位论文答辩。预计答辩时间:2027 年 6 月。",
]
empties7=[p for p in r7.paragraphs[1:] if not p.text.strip()]
for i,line in enumerate(plan):
    if i<len(empties7): fill_empty(empties7[i], line, ref7)
# 三个完成日期
repl(r7.paragraphs[-3] if False else next(p for p in r7.paragraphs if "文献阅读、科研调查计划完成日期" in p.text),
     "文献阅读、科研调查计划完成日期：", "文献阅读、科研调查计划完成日期：2026 年 9 月")
repl(next(p for p in r7.paragraphs if "论文实际工作预计完成日期" in p.text),
     "论文实际工作预计完成日期：", "论文实际工作预计完成日期：2027 年 3 月")
repl(next(p for p in r7.paragraphs if "论文撰写预计完成日期" in p.text),
     "论文撰写预计完成日期：", "论文撰写预计完成日期：2027 年 5 月")

d.save(OUT)
print("SAVED", OUT)
